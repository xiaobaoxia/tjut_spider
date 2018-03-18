# -*- coding: utf-8 -*-

# Define your item pipelines here
#
# Don't forget to add your pipeline to the ITEM_PIPELINES setting
# See: http://doc.scrapy.org/en/latest/topics/item-pipeline.html
from docx import Document
from docx.shared import Inches
import requests
import sys
import re
import StringIO
from lxml import etree
reload(sys)
sys.setdefaultencoding('utf-8')


class TjutPipeline(object):

    def process_item(self, item, spider):

        document = Document()
        document.add_heading(item['title'], 0)

        document.add_heading(item['details'], level=2)

        for i in item['content']:
            i = etree.HTML(i)
            if i.xpath('//img/@src'):
                img = StringIO.StringIO()
                img.write(requests.get('http://news.tjut.edu.cn' + i.xpath('//img/@src')[0][5:], timeout=3).content)
                img.seek(0)
                document.add_picture(img, width=Inches(6))

            else:
                document.add_paragraph(i.xpath('string(.)'))
                print i.xpath('string(.)')


        # 文档保存目录, 需要修改
        document.save('/Users/xiaobaoxia/Desktop/test/tjut/tjut/news/' + item['title'].encode('utf-8')+'.docx')

        return item
